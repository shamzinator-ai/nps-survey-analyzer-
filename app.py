import os
import json
from io import BytesIO, StringIO
import re
from typing import List, Tuple
import textwrap
import time
import asyncio
import hashlib
from langdetect import detect
from langdetect.lang_detect_exception import LangDetectException

import altair as alt
import openai
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import tempfile
import zipfile
import uuid
from utils import concat_series

# Set your OpenAI API key via environment variable
# Use an environment variable if available but also allow entering the API key
# via the UI so the app can run without preconfigured environment variables.
if "api_key" not in st.session_state:
    st.session_state["api_key"] = os.getenv("OPENAI_API_KEY", "")

openai.api_key = st.session_state["api_key"]
# Create a separate async client for concurrent API calls while keeping the
# synchronous client for single requests.
openai_async = openai.AsyncOpenAI(api_key=openai.api_key)
DEBUG = os.getenv("DEBUG", "0") == "1"

if "pdf_pivots" not in st.session_state:
    st.session_state["pdf_pivots"] = None

if not openai.api_key:
    st.sidebar.warning("Please enter your OpenAI API key to continue.")
    key_input = st.sidebar.text_input("OpenAI API Key", type="password")
    if key_input:
        openai.api_key = key_input.strip()
        st.session_state["api_key"] = openai.api_key
        openai_async = openai.AsyncOpenAI(api_key=openai.api_key)
    else:
        st.stop()
MODEL = "gpt-4o-mini"
# Approximate model cost per 1K tokens in USD used for cost estimates
TOKEN_COST_PER_1K = 0.01

# Retry configuration for API calls
MAX_RETRIES = 3
REQUEST_TIMEOUT = 20  # seconds


# Provide user-friendly messages for common OpenAI errors
def format_openai_error(e: Exception) -> str:
    if DEBUG:
        print(f"DEBUG OpenAI error: {e.__class__.__name__}: {e}", flush=True)
        if getattr(e, "__cause__", None):
            print(f"DEBUG cause: {e.__cause__}", flush=True)
    if isinstance(e, openai.AuthenticationError):
        return "Authentication failed. Check your OPENAI_API_KEY."
    if isinstance(e, openai.RateLimitError):
        return "Rate limit exceeded. Please wait and try again."
    if isinstance(e, openai.APIConnectionError):
        return "Failed to connect to OpenAI API. Check your network."
    if isinstance(e, openai.APITimeoutError):
        return "OpenAI request timed out."
    if isinstance(e, openai.APIStatusError):
        return f"OpenAI API error {e.status_code}: {e.message}"
    return str(e)


# Directory for cached processed data
CACHE_DIR = "cache"
os.makedirs(CACHE_DIR, exist_ok=True)

CATEGORIES = [
    "Search/Navigation",
    "Resource Mention",
    "User Question",
    "Translation Mention",
    "User Suggestion",
    "Pain Point",
    "AI",
    "Competitor",
    "Site Error",
    "Social Media",
    "Phonics",
    "Price Mention",
    "Accidental Purchase",
    "Resource Preview",
    "Resource Request",
    "Editing/Adapting Resource",
    "Resource Quality",
    "EDI",
    "SEND",
    "Partnership",
    "Parental Leave",
    "Email",
    "Email Verification",
    "Not Used Enough",
    "Legal",
    "Glassdoor",
    "GDPR",
    "Free Resources",
    "Download Issues",
    "Content Errors",
    "Account Access",
    "Already Cancelled",
    "Auto-renewal",
    "Book Club",
    "Cancellation Difficulty",
    "CS General",
    "CS Negative",
    "CS Positive",
    "Negative Words",
    "Positive Words",
]

# Default column names for common Twinkl survey exports
DEFAULT_USER_ID_COLUMN = "25.0: User ID"
DEFAULT_LOCATION_COLUMN = "country"
DEFAULT_NPS_COLUMN = "1: How likely are you to recommend Twinkl to a friend or colleague?"
DEFAULT_FREE_TEXT_COLUMNS = [
    "2: Thanks. We’d love to know more about why you’d recommend Twinkl.",
    "3: Thanks. Please tell us more about your score.",
    "4: Please tell us more on how we can improve your experience with Twinkl.",
    "7: We'd love to know more about your answers, especially where we can improve your website experience.",
    "10: We'd love to know more about your answers, especially where we can improve the quality of our content for you.",
    "22: How can we improve your understanding of your membership and any of the products and features you already use?",
    "23: We’d love to understand more about your answer to this question, and how we can make your subscription easier to understand.",
    "24: Is there anything else you would like to tell us about your Twinkl experience?",
]

# Columns to exclude from structured analysis
EXCLUDED_STRUCTURED_COLUMNS = [
    "Date",
    "Completed",
    "25.0: User ID",
    "26.0: t_s_id",
    "27.0: t_ca_id",
    "28.0: t_c_id",
    "29.0: t_co_id",
    "simplified sub status",
    "bundle",
    "career",
    "career_catergory",
    "country",
    "county",
    "Concatenated",
    "Translated",
    "CategoryReasoning",
    "OriginalCategories",
    "OriginalCategoryReasoning",
    "FinishReason",
    "Flagged",
    "ModelTokens",
    "10: We'd love to know more about your answers, especially where we can improve the quality of our content for you.",
]

# Predefined segment configurations used to auto-populate filters
PREDEFINED_SEGMENTS = {
    "UK Parents": {
        "location_values": ["England"],
        "filters": {"simplified career": ["Parent"]},
    },
    "US Teachers": {
        "location_values": ["United States"],
        "filters": {"simplified career": ["Teacher"]},
    },
}

# Map each category to a short description shown in the sidebar
CATEGORY_DESCRIPTIONS = {
    "Search/Navigation": "Finding resources or moving around the site",
    "Resource Mention": "References to specific products or resources",
    "User Question": "Questions users ask about the service",
    "Translation Mention": "Comments about translation quality or requests",
    "User Suggestion": "Ideas or feature requests from users",
    "Pain Point": "Descriptions of frustrations or obstacles",
    "AI": "Mentions of AI or automated features",
    "Competitor": "Comparisons to or mentions of competitors",
    "Site Error": "Reports of errors or broken pages",
    "Social Media": "Links or references to social media platforms",
    "Phonics": "Feedback specifically about phonics content",
    "Price Mention": "Concerns related to pricing or cost",
    "Accidental Purchase": "Unintended purchases or charges",
    "Resource Preview": "Ability to preview resources before buying",
    "Resource Request": "Requests for new resources",
    "Editing/Adapting Resource": "Need to edit or customise resources",
    "Resource Quality": "Opinions about quality of resources",
    "EDI": "Equity, diversity and inclusion topics",
    "SEND": "Special educational needs and disabilities",
    "Partnership": "Potential or ongoing partnerships",
    "Parental Leave": "Questions about parental leave policies",
    "Email": "General email communication issues",
    "Email Verification": "Problems verifying email accounts",
    "Not Used Enough": "Users saying they don't use the service often",
    "Legal": "Legal references or compliance concerns",
    "Glassdoor": "Mentions of Glassdoor reviews or reputation",
    "GDPR": "Data protection and GDPR related comments",
    "Free Resources": "Discussion of free offerings",
    "Download Issues": "Trouble downloading files or resources",
    "Content Errors": "Mistakes found in content or resources",
    "Account Access": "Login or account access problems",
    "Already Cancelled": "Users claiming they already cancelled",
    "Auto-renewal": "Concerns about auto-renewing subscriptions",
    "Book Club": "References to book club features",
    "Cancellation Difficulty": "Difficulty cancelling subscriptions",
    "CS General": "General customer service feedback",
    "CS Negative": "Negative comments about support",
    "CS Positive": "Positive comments about support",
    "Negative Words": "Use of negative language or sentiment",
    "Positive Words": "Use of positive language or sentiment",
}

# Map prefixes for multi-select questions to the full question text
MULTISELECT_QUESTION_TEXTS = {
    "11": ("Have you ever created your own content on the website? " "(Choose all that apply)"),
}

# Rating scale orders used for stacked bar charts and pivots
RATING_ORDER_EASE = [
    "Very Hard",
    "Hard",
    "Somewhat Hard",
    "Neutral",
    "Somewhat Easy",
    "Easy",
    "Very Easy",
]

CONTENT_RATING_ORDER = [
    "Very Poor",
    "Poor",
    "Below Average",
    "Average",
    "Above Average",
    "Good",
    "Excellent",
]

# Rating order for question 12 satisfaction levels
SATISFACTION_ORDER = [
    "Very Dissatisfied",
    "Dissatisfied",
    "Slightly Dissatisfied",
    "Neutral",
    "Indifferent",
    "Slightly Satisfied",
    "Satisfied",
    "Very Satisfied",
]

# Rating order for question 21 importance scale
IMPORTANCE_ORDER = [
    "1 - Not at all important",
    "2",
    "3",
    "4",
    "5 - Very important",
]

# Generic numeric 1-5 order used when the data contains only numbers
IMPORTANCE_ORDER_NUMERIC = ["1", "2", "3", "4", "5"]

# Color palettes for rating scales: negatives=pink, neutral=yellow, positives=blue
NEGATIVE_COLORS = ["#ffb3c6", "#ff7aa9", "#ff4d94"]
NEUTRAL_COLOR = "#ffd966"
POSITIVE_COLORS = ["#99cfff", "#66b2ff", "#3399ff"]


def rating_colors(order: List[str]) -> List[str]:
    """Return a color palette matching the given rating order."""
    if order in (RATING_ORDER_EASE, CONTENT_RATING_ORDER, SATISFACTION_ORDER):
        n_negative = 3
        n_positive = 3
        n_neutral = max(len(order) - n_negative - n_positive, 1)
        return (
            NEGATIVE_COLORS[:n_negative]
            + [NEUTRAL_COLOR] * n_neutral
            + POSITIVE_COLORS[:n_positive]
        )
    if order in (IMPORTANCE_ORDER, IMPORTANCE_ORDER_NUMERIC):
        return NEGATIVE_COLORS[1:] + [NEUTRAL_COLOR] + POSITIVE_COLORS[1:]
    # Fallback gradient
    return ["#ff66b3", "#3399ff"]


def wrap_text(text: str, width: int = 20) -> str:
    """Return text wrapped with newline characters at the given width."""
    return "\n".join(textwrap.wrap(str(text), width=width))


# ----------------------------- Utility Functions -----------------------------


def detect_language_offline(text: str) -> str:
    """Detect language locally using langdetect."""
    try:
        return detect(text)
    except LangDetectException:
        return ""


@st.cache_data(show_spinner=False)
def translate_text(text: str) -> Tuple[str, str]:
    """Detect language and translate text to English using GPT-4o-mini."""
    if not text or not text.strip():
        return "", ""
    lang = detect_language_offline(text)
    if lang.lower().startswith("en"):
        return text, "English"
    prompt = (
        "Detect the language of the following text and translate it to English. "
        "Respond in JSON with keys 'language' and 'translation'.\nText: " + text
    )
    try:
        response = openai.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            response_format={"type": "json_object"},
        )
        content = response.choices[0].message.content
        data = json.loads(content)
        return data.get("translation", "").strip(), data.get("language", "")
    except json.JSONDecodeError as e:
        st.error(f"Failed to parse JSON response: {e}")
    except Exception as e:
        if DEBUG:
            st.exception(e)
        st.error(f"Translation failed: {format_openai_error(e)}")
        return text, ""


@st.cache_data(show_spinner=False)
def categorize_text(text: str) -> List[str]:
    """Categorize text using GPT-4o-mini."""
    if not text:
        return []
    categories_str = ", ".join(CATEGORIES)
    system_prompt = (
        "You are a helpful assistant that tags survey comments with all relevant "
        "categories from the provided list. Look for any mention of AI, including "
        "references to the 'create' tool, the 'report writer' or the 'lesson "
        "planner'. Return a comma-separated list of matching categories. If none "
        "apply, return 'None'."
    )
    user_prompt = f"Categories: {categories_str}\nComment: {text}"
    try:
        response = openai.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0,
            response_format={"type": "json_object"},
        )
        result = response.choices[0].message.content.strip()
        data = json.loads(result)
        categories_raw = data.get("categories", "")
        if categories_raw.lower() == "none":
            return []
        return [c.strip() for c in categories_raw.split(",")]
    except json.JSONDecodeError as e:
        st.error(f"Failed to parse JSON response: {e}")
        return []
    except Exception as e:
        if DEBUG:
            st.exception(e)
        st.error(f"Categorization failed: {format_openai_error(e)}")
        return []


async def async_translate_batch(texts: List[str]) -> List[Tuple[str, str, int, str]]:
    """Translate a batch of texts concurrently using OpenAI's async API."""

    async def _translate(text: str) -> Tuple[str, str, int, str]:
        if not text or not text.strip():
            return "", "", 0, ""
        lang = detect_language_offline(text)
        if lang.lower().startswith("en"):
            return text, "English", 0, "skipped"
        prompt = (
            "Detect the language of the following text and translate it to English."
            "Respond in JSON with keys 'language' and 'translation'.\nText: " + text
        )
        delay = 1
        for attempt in range(MAX_RETRIES):
            try:
                response = await asyncio.wait_for(
                    openai_async.chat.completions.create(
                        model=MODEL,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0,
                        response_format={"type": "json_object"},
                    ),
                    timeout=REQUEST_TIMEOUT,
                )
                content = response.choices[0].message.content
                data = json.loads(content)
                tokens = response.usage.total_tokens if getattr(response, "usage", None) else 0
                finish = response.choices[0].finish_reason or ""
                return (
                    data.get("translation", "").strip(),
                    data.get("language", ""),
                    tokens,
                    finish,
                )
            except (json.JSONDecodeError, Exception) as e:
                if attempt < MAX_RETRIES - 1:
                    await asyncio.sleep(delay)
                    delay *= 2
                    continue
                if isinstance(e, json.JSONDecodeError):
                    st.error(f"Failed to parse JSON response: {e}")
                else:
                    if DEBUG:
                        st.exception(e)
                    st.warning(f"Translation failed after retries: {format_openai_error(e)}")
                return text, "", 0, "error"

    tasks = [asyncio.create_task(_translate(t)) for t in texts]
    return await asyncio.gather(*tasks)


async def async_categorize_batch(texts: List[str]) -> List[Tuple[List[str], str, int, str]]:
    """Categorize a batch of texts concurrently with short reasoning."""

    categories_str = ", ".join(CATEGORIES)
    system_prompt = (
        "You are a helpful assistant that tags survey comments with all relevant "
        "categories from the provided list. Look for any mention of AI, including "
        "references to the 'create' tool, the 'report writer' or the 'lesson "
        "planner'. Return a JSON object with keys 'categories' (comma-separated "
        "list of categories or 'None') and 'reasoning' (max 30 words explaining "
        "the choice)."
    )

    async def _categorize(text: str) -> Tuple[List[str], str, int, str]:
        if not text:
            return [], "", 0, ""
        user_prompt = f"Categories: {categories_str}\nComment: {text}"
        delay = 1
        for attempt in range(MAX_RETRIES):
            try:
                response = await asyncio.wait_for(
                    openai_async.chat.completions.create(
                        model=MODEL,
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_prompt},
                        ],
                        temperature=0,
                        response_format={"type": "json_object"},
                    ),
                    timeout=REQUEST_TIMEOUT,
                )
                content = response.choices[0].message.content.strip()
                data = json.loads(content)
                categories_raw = data.get("categories", "")
                reasoning = data.get("reasoning", "").strip()
                tokens = response.usage.total_tokens if getattr(response, "usage", None) else 0
                finish = response.choices[0].finish_reason or ""
                if categories_raw.lower() == "none":
                    return [], reasoning, tokens, finish
                categories = [c.strip() for c in categories_raw.split(",")]
                return categories, reasoning, tokens, finish
            except (json.JSONDecodeError, Exception) as e:
                if attempt < MAX_RETRIES - 1:
                    await asyncio.sleep(delay)
                    delay *= 2
                    continue
                if isinstance(e, json.JSONDecodeError):
                    st.error(f"Failed to parse JSON response: {e}")
                else:
                    if DEBUG:
                        st.exception(e)
                    st.warning(f"Categorization failed after retries: {format_openai_error(e)}")
                return [], "", 0, "error"

    tasks = [asyncio.create_task(_categorize(t)) for t in texts]
    return await asyncio.gather(*tasks)


def generate_pivot(
    df: pd.DataFrame, column: str, *, exclude_na: bool = False
) -> pd.DataFrame:
    """Return value counts with percentage and total row.

    Parameters
    ----------
    df : pd.DataFrame
        DataFrame containing the survey data.
    column : str
        Column to summarize.
    exclude_na : bool, optional
        When True, ignore missing values so ``NaN`` does not appear in charts.
    """
    series = df[column]
    if exclude_na:
        series = series.dropna()
    pivot = series.value_counts(dropna=False).reset_index()
    pivot.columns = ["Response", "Count"]
    total = pivot["Count"].sum()
    pivot["Percent"] = (pivot["Count"] / total * 100).round(1)
    total_row = pd.DataFrame({"Response": ["Total"], "Count": [total], "Percent": [100.0]})
    pivot = pd.concat([pivot, total_row], ignore_index=True)
    return pivot


def multiselect_pivot(df: pd.DataFrame, columns: List[str]) -> pd.DataFrame:
    """Aggregate binary multi-select columns into a single pivot."""
    rows = []
    for col in columns:
        label = re.sub(r"^\d+[\.:]\s*", "", col).strip()
        values = df[col].fillna("")
        yes_vals = {"1", "1.0", "true", "yes", "checked"}
        count = sum(str(v).strip().lower() in yes_vals for v in values)
        rows.append({"Response": label, "Count": int(count)})
    pivot = pd.DataFrame(rows)
    total = pivot["Count"].sum()
    pivot["Percent"] = (pivot["Count"] / total * 100).round(1) if total else 0
    total_row = pd.DataFrame(
        {"Response": ["Total"], "Count": [total], "Percent": [100.0 if total else 0.0]}
    )
    return pd.concat([pivot, total_row], ignore_index=True)


def rating_pivot(
    df: pd.DataFrame, columns: List[str], order: List[str] | None = None
) -> pd.DataFrame:
    """Combine multiple rating columns into a single long-format pivot."""
    if order is None:
        order = RATING_ORDER_EASE
    df_long = df[columns].melt(value_name="Rating", var_name="Aspect").dropna()
    df_long["Rating"] = (
        df_long["Rating"].apply(lambda x: str(x).strip()).str.replace(r"\.0$", "", regex=True)
    )
    df_long["Aspect"] = df_long["Aspect"].apply(
        lambda c: re.sub(r"^\d+[\.:]\s*", "", str(c)).split(":")[-1].strip()
    )
    pivot = df_long.value_counts(["Aspect", "Rating"]).reset_index(name="Count")
    pivot["Percent"] = pivot.groupby("Aspect")["Count"].transform(
        lambda x: (x / x.sum() * 100).round(1)
    )
    pivot["Rating"] = pd.Categorical(pivot["Rating"], categories=order, ordered=True)
    return pivot.sort_values(["Aspect", "Rating"]).reset_index(drop=True)


def combined_rating_pivot(
    df: pd.DataFrame, columns: List[str], order: List[str] | None = None
) -> pd.DataFrame:
    """Aggregate rating columns into a single pivot table."""
    if not columns:
        return pd.DataFrame(columns=["Response", "Count", "Percent"])
    if order is None:
        order = SATISFACTION_ORDER
    ratings = (
        df[columns]
        .melt(value_name="Response")["Response"]
        .dropna()
        .apply(lambda x: str(x).strip())
        .str.replace(r"\.0$", "", regex=True)
    )
    pivot = ratings.value_counts().reset_index()
    pivot.columns = ["Response", "Count"]
    total = pivot["Count"].sum()
    pivot["Percent"] = (pivot["Count"] / total * 100).round(1)
    pivot["Response"] = pd.Categorical(pivot["Response"], categories=order, ordered=True)
    pivot = pivot.sort_values("Response").reset_index(drop=True)
    total_row = pd.DataFrame({"Response": ["Total"], "Count": [total], "Percent": [100.0]})
    return pd.concat([pivot, total_row], ignore_index=True)


def stacked_bar_chart(pivot: pd.DataFrame, title: str, order: List[str] | None = None) -> BytesIO:
    """Display a stacked bar chart and allow PNG/SVG download."""
    if order is None:
        order = RATING_ORDER_EASE

    # Pre-wrap aspect labels to avoid JavaScript expressions which break during
    # PNG conversion. Insert newline characters to wrap long text.
    pivot = pivot.copy()
    if "Aspect" in pivot.columns:
        pivot["Aspect_wrapped"] = pivot["Aspect"].apply(wrap_text)

    chart = (
        alt.Chart(pivot, background="white")
        .mark_bar()
        .encode(
            x=alt.X(
                "Aspect_wrapped:N",
                title="Aspect",
                axis=alt.Axis(labelAngle=-90, labelLimit=0),
            ),
            y=alt.Y(
                "Count:Q",
                stack="normalize",
                axis=alt.Axis(format="%"),
                title="Percent",
            ),
            color=alt.Color(
                "Rating:N",
                sort=order,
                scale=alt.Scale(domain=order, range=rating_colors(order)),
            ),
            tooltip=["Rating", "Count"],
        )
        .properties(title=f"{title} \U0001F4CA", height=300)
        .configure_title(fontSize=22)
        .configure_axis(labelFontSize=16, titleFontSize=18)
    )
    st.altair_chart(chart, use_container_width=True)

    png_buffer = BytesIO()
    chart.save(png_buffer, format="png")
    png_buffer.seek(0)
    svg_buffer = StringIO()
    chart.save(svg_buffer, format="svg")
    svg_bytes = svg_buffer.getvalue().encode("utf-8")
    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "Download Chart PNG",
            png_buffer,
            file_name=f"{title}.png",
            mime="image/png",
            key=unique_key(f"{title}_png"),
        )
    with col2:
        st.download_button(
            "Download Chart SVG",
            svg_bytes,
            file_name=f"{title}.svg",
            mime="image/svg+xml",
            key=unique_key(f"{title}_svg"),
        )
    return png_buffer


def create_chart(pivot: pd.DataFrame, title: str, order: List[str] | None = None):
    """Return an Altair chart object matching the on-screen chart."""
    color_range = ["#ff66b3", "#3399ff"]

    # Ensure the expected column name exists for the X encoding.
    pivot = pivot.rename(columns={"Category": "Response"})
    if "Response" in pivot.columns:
        # Ensure we work with strings for label wrapping and comparisons
        pivot["Response"] = pivot["Response"].astype(str)
        # Remove any total row so it doesn't appear as its own bar
        pivot = pivot[pivot["Response"] != "Total"]

        # Pre-wrap labels to avoid using JavaScript expressions which can fail
        # during PNG conversion in vl-convert. Wrap long text with newlines.
        pivot["Response_wrapped"] = pivot["Response"].astype(str).apply(wrap_text)

    enc_color = alt.Color(
        "Count:Q",
        scale=alt.Scale(range=color_range),
        legend=None,
    )
    if order:
        enc_color = alt.Color(
            "Response:N",
            sort=order,
            scale=alt.Scale(domain=order, range=rating_colors(order)),
        )

    chart = (
        alt.Chart(pivot, background="white")
        .mark_bar()
        .encode(
            x=alt.X(
                "Response_wrapped:N",
                sort="-y",
                title="Response",
                axis=alt.Axis(labelAngle=0, labelLimit=0),
            ),
            y=alt.Y("Count:Q", title="Count"),
            color=enc_color,
            tooltip=["Response", "Count"],
        )
        .properties(title=f"{title} 📊", height=300)
        .configure_title(fontSize=22)
        .configure_axis(labelFontSize=16, titleFontSize=18)
    )
    return chart


def chart_png(pivot: pd.DataFrame, title: str, order: List[str] | None = None) -> BytesIO:
    """Return a PNG image buffer for the given pivot chart."""
    chart = create_chart(pivot, title, order=order)
    buf = BytesIO()
    chart.save(buf, format="png")
    buf.seek(0)
    return buf


def bar_chart(pivot: pd.DataFrame, title: str, order: List[str] | None = None) -> BytesIO:
    """Display a bar chart and provide PNG/SVG downloads. Returns PNG buffer."""
    chart = create_chart(pivot, title, order=order)
    st.altair_chart(chart, use_container_width=True)

    png_buffer = BytesIO()
    chart.save(png_buffer, format="png")
    png_buffer.seek(0)
    svg_buffer = StringIO()
    chart.save(svg_buffer, format="svg")
    svg_bytes = svg_buffer.getvalue().encode("utf-8")
    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "Download Chart PNG",
            png_buffer,
            file_name=f"{title}.png",
            mime="image/png",
            key=unique_key(f"{title}_png"),
        )
    with col2:
        st.download_button(
            "Download Chart SVG",
            svg_bytes,
            file_name=f"{title}.svg",
            mime="image/svg+xml",
            key=unique_key(f"{title}_svg"),
        )

    return png_buffer


def safe_name(name: str) -> str:
    """Return a filesystem-friendly version of a name."""
    name = name.strip().replace(" ", "_")
    return re.sub(r"[^A-Za-z0-9_-]", "_", name)


def unique_key(name: str) -> str:
    """Return a unique key for Streamlit widgets based on the provided name."""
    return f"{safe_name(name)}_{uuid.uuid4().hex}"


def category_frequency(df: pd.DataFrame) -> pd.DataFrame:
    """Return counts and percentages for all categories including blanks."""
    if "Categories" not in df.columns:
        return pd.DataFrame(columns=["Category", "Count", "Percent"])

    cat_lists: list[str] = []
    for cats in df["Categories"].fillna(""):
        parts = [c.strip() for c in str(cats).split(",") if c.strip()]
        if not parts:
            parts = ["Uncategorized"]
        cat_lists.extend(parts)

    pivot = pd.Series(cat_lists).value_counts().reset_index()
    pivot.columns = ["Category", "Count"]
    total = pivot["Count"].sum()
    pivot["Percent"] = (pivot["Count"] / total * 100).round(1)
    total_row = pd.DataFrame({"Category": ["Total"], "Count": [total], "Percent": [100.0]})
    return pd.concat([pivot, total_row], ignore_index=True)


def sentiment_metrics(df: pd.DataFrame) -> tuple[int, int]:
    """Return counts of positive and negative comments."""
    if "Categories" not in df.columns:
        return 0, 0
    pos = df["Categories"].str.contains("Positive Words", na=False).sum()
    neg = df["Categories"].str.contains("Negative Words", na=False).sum()
    return pos, neg


def compute_nps_score(df: pd.DataFrame, column: str) -> int | None:
    """Return the Net Promoter Score for a numeric 0-10 column."""
    if column not in df.columns:
        return None
    values = pd.to_numeric(df[column], errors="coerce").dropna()
    if values.empty:
        return None
    promoters = (values >= 9).sum()
    detractors = (values <= 6).sum()
    total = len(values)
    return int(round((promoters - detractors) / total * 100))


def compute_kpis(
    df: pd.DataFrame, nps_col: str | None
) -> tuple[pd.DataFrame, pd.DataFrame, tuple[int, int], int | None]:
    """Calculate NPS metrics, category frequency and sentiment metrics."""
    nps_pivot = pd.DataFrame()
    nps_score = None
    if nps_col and nps_col in df.columns:
        nps_pivot = generate_pivot(df, nps_col)
        nps_score = compute_nps_score(df, nps_col)
    cat_pivot = category_frequency(df)
    sentiment = sentiment_metrics(df)
    return nps_pivot, cat_pivot, sentiment, nps_score


def display_summary(df: pd.DataFrame, nps_col: str | None):
    """Show high-level KPIs and charts."""
    st.subheader("🚀 High-Level KPIs")
    st.metric("Rows After Filters", len(df))
    total_tokens = st.session_state.get("total_tokens")
    if total_tokens is None and "ModelTokens" in df.columns:
        total_tokens = int(df["ModelTokens"].sum())
    if total_tokens is not None:
        st.metric("Total Model Tokens", total_tokens)
        if TOKEN_COST_PER_1K:
            est_cost = total_tokens / 1000 * TOKEN_COST_PER_1K
            st.metric("Estimated Cost", f"${est_cost:.2f}")
    nps_pivot, cat_pivot, (pos, neg), nps_score = compute_kpis(df, nps_col)
    if nps_score is not None:
        st.metric("NPS Score", nps_score)
    if not nps_pivot.empty:
        st.write("### NPS Distribution")
        st.dataframe(nps_pivot)
        bar_chart(nps_pivot, "NPS Distribution")
    if not cat_pivot.empty:
        st.write("### Category Frequency")
        st.dataframe(cat_pivot)
        cat_chart = cat_pivot[cat_pivot["Category"] != "Uncategorized"]
        if not cat_chart.empty:
            bar_chart(cat_chart, "Category Frequency")
        st.metric("Positive/Negative Ratio", f"{pos}:{neg}")
        st.write("Top 3 Issues:", ", ".join(cat_pivot.head(3)["Category"].tolist()))


def download_link(df: pd.DataFrame, filename: str, label: str, help: str | None = None):
    csv = df.to_csv(index=False).encode("utf-8")
    st.download_button(
        label,
        csv,
        file_name=filename,
        mime="text/csv",
        help=help,
        key=unique_key(filename),
    )


def export_excel(df: pd.DataFrame, filename: str, label: str, help: str | None = None):
    """Download a DataFrame as an Excel file."""
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False)
    st.download_button(
        label,
        buffer.getvalue(),
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        help=help,
        key=unique_key(filename),
    )


def export_full_excel(df: pd.DataFrame, filename: str, label: str, help: str | None = None):
    """Download the entire processed DataFrame as an Excel file."""
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False)
    st.download_button(
        label,
        buffer.getvalue(),
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        help=help,
        key=unique_key(filename),
    )


def read_uploaded_file(uploaded_file) -> pd.DataFrame | None:
    """Load an uploaded CSV or Excel file with encoding validation."""
    raw_bytes = uploaded_file.getvalue()
    try:
        if uploaded_file.name.endswith(("xls", "xlsx")):
            return pd.read_excel(BytesIO(raw_bytes))
        try:
            return pd.read_csv(BytesIO(raw_bytes), encoding="utf-8")
        except UnicodeDecodeError:
            st.warning("File is not UTF-8 encoded, attempting latin-1 decoding")
            return pd.read_csv(BytesIO(raw_bytes), encoding="latin-1")
    except UnicodeDecodeError as e:
        st.error(f"Encoding error: {e}")
    except Exception as e:
        st.error(f"Failed to read file: {e}")
    return None


def validate_file(uploaded_file) -> bool:
    """Basic checks for uploaded file size and type."""
    max_size = 10 * 1024 * 1024  # 10MB
    if uploaded_file.size > max_size:
        st.error("File too large. Maximum size is 10MB.")
        return False
    return True


def validate_columns(
    user_id_col: str, location_col: str, free_text_cols: List[str], structured_cols: List[str]
) -> bool:
    """Ensure mandatory columns are selected."""
    errors = []
    if not user_id_col:
        errors.append("Please select a user ID column.")
    if not location_col:
        errors.append("Please select a location column.")
    if not free_text_cols:
        errors.append("Select at least one free-text column.")
    if not structured_cols:
        errors.append("Select at least one structured column.")
    if errors:
        for msg in errors:
            st.error(msg)
        return False
    return True


def review_translations(df: pd.DataFrame, id_col: str) -> pd.DataFrame:
    """Allow user to edit translations and categories."""
    show_reasoning = st.checkbox(
        "Show category reasoning",
        value=False,
        help="Display the explanation behind each auto-assigned category.",
    )
    flags = []
    for idx, row in df.iterrows():
        user_val = row.get(id_col, "Unknown")
        with st.expander(f"User {user_val}"):
            st.write("**Original:**", row["Concatenated"])
            st.write(f"Tokens used: {row.get('ModelTokens', 0)}")
            st.write(f"Finish reason: {row.get('FinishReason', '')}")
            if show_reasoning and row.get("OriginalCategoryReasoning"):
                st.write(
                    "**Original Reasoning:**",
                    row["OriginalCategoryReasoning"],
                )
            if show_reasoning and row.get("CategoryReasoning"):
                st.write("**Translated Reasoning:**", row["CategoryReasoning"])
            new_trans = st.text_area(
                "Translated",
                value=row["Translated"],
                key=f"trans_{idx}",
                help="Edit the AI translation if it looks incorrect.",
            )
            new_orig_cats = st.multiselect(
                "Original Categories",
                options=CATEGORIES,
                default=[
                    c.strip()
                    for c in row.get("OriginalCategories", "").split(",")
                    if c
                ],
                key=f"orig_cat_{idx}",
                help="Edit categories for the original text.",
            )
            new_cats = st.multiselect(
                "Translated Categories",
                options=CATEGORIES,
                default=[c.strip() for c in row["Categories"].split(",") if c],
                key=f"cat_{idx}",
                help="Add or remove categories for this comment.",
            )
            flag = st.checkbox(
                "Flag for review", key=f"flag_{idx}", help="Mark this comment for manual follow-up."
            )
            df.at[idx, "Translated"] = new_trans
            df.at[idx, "OriginalCategories"] = ", ".join(new_orig_cats)
            df.at[idx, "Categories"] = ", ".join(new_cats)
            flags.append(flag)
    df["Flagged"] = flags
    return df


def generate_report(df: pd.DataFrame) -> str:
    """Generate a detailed report covering all requested sections."""
    prompt = (
        "You are an expert analyst summarising NPS survey feedback. "
        "Write a detailed report with the following sections:\n"
        "Executive Summary of findings;\n"
        "Customer Suggestions with quoted comments and IDs;\n"
        "Localisation Issues, especially for US users;\n"
        "Pain Points;\n"
        "Negative Comments;\n"
        "Feedback not captured in categories;\n"
        "Suggested Next Steps with rationale, priority and confidence level;\n"
        "Devils Advocate (limitations and possible biases);\n"
        "Conclusion with summary and non-obvious recommendations.\n"
        "Every assertion must be supported by one to ten verbatim quotes "
        "from the provided comments including User IDs. "
        "For every statement include raw number and percentage of supporting comments. "
        "Do not invent data; reference only the provided input. "
        "Avoid bullet points and write in narrative paragraphs."
    )
    user_content = df.to_csv(index=False)
    try:
        response = openai.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": user_content},
            ],
            temperature=0.3,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        if DEBUG:
            st.exception(e)
        st.error(f"Report generation failed: {format_openai_error(e)}")
        return ""


def save_docx(text: str, pivots: dict[str, pd.DataFrame]) -> BytesIO:
    """Create a DOCX report with text, pivot tables and charts."""
    document = Document()
    for para in text.split("\n"):
        document.add_paragraph(para)

    for title, pivot in pivots.items():
        document.add_heading(title, level=2)
        table = document.add_table(rows=1, cols=len(pivot.columns))
        hdr = table.rows[0].cells
        for idx, col in enumerate(pivot.columns):
            hdr[idx].text = str(col)
        for _, row in pivot.iterrows():
            cells = table.add_row().cells
            for idx, col in enumerate(pivot.columns):
                cells[idx].text = str(row[col])
        img = chart_png(pivot, f"{title} Responses")
        img.name = "chart.png"
        document.add_picture(img, width=Inches(6))

    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio


def save_pdf(
    text: str,
    pivots: dict[str, pd.DataFrame],
    include_charts: bool = True,
    include_tables: bool = True,
) -> BytesIO:
    """Save text, charts and tables as a PDF.

    Parameters
    ----------
    text : str
        Introductory text to display at the top of the PDF.
    pivots : dict[str, pd.DataFrame]
        Mapping of question titles to pivot DataFrames.
    include_charts : bool, optional
        Include chart images under each question, by default ``True``.
    include_tables : bool, optional
        Include pivot tables under each question, by default ``True``.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in text.split("\n"):
        pdf.multi_cell(0, 10, line)

    for title, pivot in pivots.items():
        pdf.ln(5)
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 10, title, ln=True)
        if include_tables:
            pdf.set_font("Arial", size=10)
            col_widths = [80, 30, 30]
            headers = list(pivot.columns)
            for i, h in enumerate(headers):
                w = col_widths[i] if i < len(col_widths) else 30
                pdf.cell(w, 8, str(h), border=1)
            pdf.ln()
            for _, row in pivot.iterrows():
                for i, h in enumerate(headers):
                    w = col_widths[i] if i < len(col_widths) else 30
                    pdf.cell(w, 8, str(row[h]), border=1)
                pdf.ln()
        if include_charts:
            img = chart_png(pivot, f"{title} Responses")
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(img.getvalue())
                tmp.flush()
                pdf.image(tmp.name, w=180)
            os.unlink(tmp.name)

    # Generate PDF bytes and write them to a BytesIO buffer
    pdf_bytes = pdf.output(dest="S").encode("latin-1")
    bio = BytesIO()
    bio.write(pdf_bytes)
    bio.seek(0)
    return bio


def process_free_text(
    df: pd.DataFrame,
    free_text_cols: List[str],
    cache_path: str,
    batch_size: int = 5,
) -> pd.DataFrame:
    """Concatenate, translate and categorize free-text columns with progress.

    The function saves intermediate results to ``cache_path`` with a ``_partial``
    suffix after each processed batch so that work can be resumed if the app is
    restarted.

    Parameters
    ----------
    batch_size : int, optional
        Number of rows to process concurrently in each batch.
    """

    df["Concatenated"] = concat_series(df, free_text_cols)

    # Ensure output columns exist
    for col, default in [
        ("Translated", ""),
        ("Language", ""),
        ("Categories", ""),
        ("CategoryReasoning", ""),
        ("OriginalCategories", ""),
        ("OriginalCategoryReasoning", ""),
        ("ModelTokens", 0),
        ("FinishReason", ""),
    ]:
        if col not in df.columns:
            df[col] = default

    # Determine which rows still need processing
    to_process = [
        idx
        for idx in df.index
        if not isinstance(df.at[idx, "Translated"], str)
        or not str(df.at[idx, "Translated"]).strip()
    ]

    progress = st.progress(0.0, text="Starting...")
    start_time = time.time()
    partial_path = cache_path.replace(".pkl", "_partial.pkl")

    for batch_start in range(0, len(to_process), batch_size):
        batch_indices = to_process[batch_start : batch_start + batch_size]
        batch_texts = [df.at[idx, "Concatenated"] for idx in batch_indices]

        orig_cats_data = asyncio.run(async_categorize_batch(batch_texts))
        batch_orig_cats = [cats for cats, _, _, _ in orig_cats_data]
        batch_orig_reason = [reason for _, reason, _, _ in orig_cats_data]
        batch_toks_orig = [tok for _, _, tok, _ in orig_cats_data]
        batch_finish_orig = [fin for _, _, _, fin in orig_cats_data]

        trans_lang = asyncio.run(async_translate_batch(batch_texts))
        batch_trans = [t for t, _, _, _ in trans_lang]
        batch_langs = [lang for _, lang, _, _ in trans_lang]
        batch_toks_trans = [tok for _, _, tok, _ in trans_lang]
        batch_finish_trans = [fin for _, _, _, fin in trans_lang]

        batch_cats_data = asyncio.run(async_categorize_batch(batch_trans))
        batch_cats = [cats for cats, _, _, _ in batch_cats_data]
        batch_reason = [reason for _, reason, _, _ in batch_cats_data]
        batch_toks_cat = [tok for _, _, tok, _ in batch_cats_data]
        batch_finish_cat = [fin for _, _, _, fin in batch_cats_data]

        for offset, idx in enumerate(batch_indices):
            df.at[idx, "Translated"] = batch_trans[offset]
            df.at[idx, "Language"] = batch_langs[offset]
            df.at[idx, "Categories"] = ", ".join(batch_cats[offset])
            df.at[idx, "CategoryReasoning"] = batch_reason[offset]
            df.at[idx, "OriginalCategories"] = ", ".join(batch_orig_cats[offset])
            df.at[idx, "OriginalCategoryReasoning"] = batch_orig_reason[offset]
            total_tokens = (
                batch_toks_trans[offset]
                + batch_toks_cat[offset]
                + batch_toks_orig[offset]
            )
            df.at[idx, "ModelTokens"] = total_tokens
            fin_orig = batch_finish_orig[offset]
            fin_trans = batch_finish_trans[offset]
            fin_cat = batch_finish_cat[offset]
            if fin_orig == fin_trans == fin_cat:
                df.at[idx, "FinishReason"] = fin_trans
            else:
                df.at[idx, "FinishReason"] = (
                    f"O:{fin_orig}; T:{fin_trans}; C:{fin_cat}"
                )

        processed = batch_start + len(batch_indices)
        rate = (time.time() - start_time) / (processed if processed else 1)
        remaining = rate * (len(to_process) - processed)
        progress.progress(processed / len(to_process), text=f"Processing... ETA {int(remaining)}s")

        # Persist partial results after each batch
        df.to_pickle(partial_path)

    progress.empty()
    return df


# ----------------------------- Streamlit App -----------------------------

st.set_page_config(page_title="NPS Survey Analyzer", page_icon="📊", layout="wide")


def apply_style():
    """Inject a modern look and feel using custom CSS."""
    st.markdown(
        """
        <link href="https://fonts.googleapis.com/css2?family=Roboto:wght@400;700&display=swap" rel="stylesheet">
        <style>
            html, body, [class*="css"]  {
                font-family: 'Roboto', sans-serif;
                color: #000;
            }
            [data-testid="stAppViewContainer"] {
                background: linear-gradient(135deg, #f4f8fb 0%, #e0ecff 100%);
                height: 100vh;
                overflow-y: auto;
            }
            .stButton>button {
                background-color: #0E79B2;
                color: white;
                border-radius: 4px;
                padding: 0.5rem 1rem;
                font-size: 1rem;
            }
            .stButton>button:hover {
                background-color: #0b6391;
                color: white;
            }
            .stProgress>div>div>div {
                background-color: #0E79B2;
            }
            @media (prefers-color-scheme: dark) {
                html, body, [class*="css"] { color: #000 !important; }
                .stButton>button { color: #fff !important; }
            }
        </style>
        """,
        unsafe_allow_html=True,
    )


apply_style()

# Sidebar toggle to switch to black text and blue buttons
if "contrast_mode" not in st.session_state:
    st.session_state["contrast_mode"] = True

contrast_toggle = st.sidebar.toggle(
    "Black text and blue buttons",
    value=st.session_state["contrast_mode"],
)
st.session_state["contrast_mode"] = contrast_toggle

if contrast_toggle:
    st.markdown(
        """
        <style id="contrast-style">
            html, body, [class*="css"] { color: #000 !important; }
            .stButton>button {
                background-color: #0E79B2 !important;
                color: #fff !important;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )
else:
    st.markdown(
        """
        <style id="contrast-style"></style>
        """,
        unsafe_allow_html=True,
    )
st.title("NPS Survey Analyzer")

st.sidebar.header("1. Upload Survey Data")
st.sidebar.markdown("🔗 [Troubleshooting guide](README.md#troubleshooting)")
file = st.sidebar.file_uploader(
    "Upload CSV, XLS or XLSX",
    type=["csv", "xls", "xlsx"],
    key="data_file",
    help="File must include a unique ID, location, at least one structured and one free-text column.",
)

if file is not None:
    st.session_state["uploaded_file"] = file
else:
    file = st.session_state.get("uploaded_file")

if file and validate_file(file):
    raw_bytes = file.getvalue()
    # Use SHA-256 rather than MD5 to generate a cache key for the uploaded file
    checksum = hashlib.sha256(raw_bytes).hexdigest()
    cache_path = os.path.join(CACHE_DIR, f"{checksum}.pkl")
    partial_path = os.path.join(CACHE_DIR, f"{checksum}_partial.pkl")

    if st.sidebar.button("Clear Cached Data"):
        if os.path.exists(cache_path):
            os.remove(cache_path)
        if os.path.exists(partial_path):
            os.remove(partial_path)
        st.session_state.pop("processed_df", None)
        st.sidebar.success("Cached data cleared")

    if "processed_df" in st.session_state:
        df = st.session_state["processed_df"]
    elif os.path.exists(cache_path):
        df = pd.read_pickle(cache_path)
        st.success("Loaded cached processed data")
    elif os.path.exists(partial_path):
        df = pd.read_pickle(partial_path)
        st.info("Resuming from saved progress")
    else:
        df = read_uploaded_file(file)
        if df is None:
            st.stop()
        if df.empty:
            st.error("Uploaded file contains no data")
            st.stop()
        if df.isnull().all(axis=0).any():
            st.warning("Some columns contain only missing values")
        if df.isnull().all(axis=1).any():
            st.warning("Some rows contain only missing values")

    if "original_df" not in st.session_state:
        st.session_state["original_df"] = df.copy()

    st.subheader("Data Preview")
    st.dataframe(df.head(10))
    st.write(f"**Rows:** {df.shape[0]}  **Columns:** {df.shape[1]}")
    st.write("**Columns:**", ", ".join(df.columns))

    user_id_default = (
        df.columns.get_loc(DEFAULT_USER_ID_COLUMN) if DEFAULT_USER_ID_COLUMN in df.columns else 0
    )
    user_id_col = st.selectbox(
        "Column with User ID",
        options=df.columns,
        index=user_id_default,
        help="Select the column that uniquely identifies each user.",
    )
    location_default = (
        df.columns.get_loc(DEFAULT_LOCATION_COLUMN) if DEFAULT_LOCATION_COLUMN in df.columns else 0
    )
    location_col = st.selectbox(
        "Column with County/Location",
        options=df.columns,
        index=location_default,
        help="Pick the column that indicates user location or segment.",
    )

    ft_options = [c for c in df.columns if c not in [user_id_col, location_col]]
    ft_default = [c for c in DEFAULT_FREE_TEXT_COLUMNS if c in ft_options]
    free_text_cols = st.multiselect(
        "Free-text response columns",
        options=ft_options,
        default=ft_default,
        help="These comments will be translated and categorised.",
    )

    available_structured = [
        c
        for c in df.columns
        if c not in free_text_cols + [user_id_col, location_col]
        and c not in EXCLUDED_STRUCTURED_COLUMNS
        and "unnamed" not in str(c).lower()
        and not str(c).startswith("5.")
        and not str(c).startswith("5:")
    ]
    structured_cols = st.multiselect(
        "Structured question columns",
        options=available_structured,
        default=available_structured,
        help="Responses to these columns will be summarised in pivot tables.",
    )

    # Skip any columns accidentally selected that contain 'Unnamed'
    structured_cols = [
        c
        for c in structured_cols
        if "unnamed" not in str(c).lower()
        and not str(c).startswith("5.")
        and not str(c).startswith("5:")
    ]

    preset_name = st.selectbox(
        "Predefined segment (optional)",
        ["None"] + list(PREDEFINED_SEGMENTS.keys()),
        help="Automatically populate filters for common segments.",
    )
    preset = PREDEFINED_SEGMENTS.get(preset_name)

    segment_options = df[location_col].dropna().unique().tolist()
    preset_segments = []
    if preset:
        preset_segments = [s for s in preset.get("location_values", []) if s in segment_options]
    selected_segments = st.multiselect(
        "Filter by segment (optional)",
        options=segment_options,
        default=preset_segments,
        help="Choose segments to focus on or leave empty for all.",
    )

    filter_col_options = [
        c for c in df.columns if c not in free_text_cols + [user_id_col, location_col]
    ]
    preset_filter_cols = []
    if preset:
        preset_filter_cols = [c for c in preset.get("filters", {}) if c in filter_col_options]
    addl_filter_cols = st.multiselect(
        "Additional segment columns to filter (optional)",
        options=filter_col_options,
        default=preset_filter_cols,
        help="Select other columns like career type to filter by.",
    )

    addl_filters = {}
    for col in addl_filter_cols:
        opts = df[col].dropna().unique().tolist()
        default_vals = []
        if preset and col in preset.get("filters", {}):
            default_vals = [v for v in preset["filters"][col] if v in opts]
        selected = st.multiselect(
            f"Values for {col}",
            options=opts,
            default=default_vals,
            key=f"values_{col}",
        )
        if selected:
            addl_filters[col] = selected

    with st.sidebar.expander("Category descriptions"):
        for cat, desc in CATEGORY_DESCRIPTIONS.items():
            st.write(f"**{cat}** - {desc}")

    analysis_mode = st.radio(
        "Analysis mode",
        ["Structured Data Only", "Free Text Only", "Both"],
        index=2,
        help="Choose which analysis steps to run",
    )

    batch_size = st.sidebar.slider(
        "Batch size",
        min_value=1,
        max_value=20,
        value=5,
        help="Number of comments to process at once",
    )

    if st.session_state.get("pdf_pivots"):
        pdf_buf_charts = save_pdf(
            "NPS Survey Charts", st.session_state["pdf_pivots"],
            include_charts=True, include_tables=False
        )
        st.download_button(
            "Download Charts PDF",
            pdf_buf_charts,
            "charts_summary.pdf",
            help="Download a PDF containing all charts with question text.",
            key=unique_key("charts_pdf_summary"),
        )
        pdf_buf_tables = save_pdf(
            "NPS Survey Tables", st.session_state["pdf_pivots"], include_charts=False
        )
        st.download_button(
            "Download Tables PDF",
            pdf_buf_tables,
            "tables_summary.pdf",
            help="Download a PDF with each pivot table and its question.",
            key=unique_key("tables_pdf_summary"),
        )

    show_comments = st.checkbox(
        "Show detailed comments",
        value=False,
        help="Display each comment with its categories for manual review.",
    )

    process_clicked = st.button(
        "Process Data", help="Translate comments, categorise them and generate summaries."
    )
    if process_clicked:
        if not validate_columns(
            user_id_col,
            location_col,
            free_text_cols,
            structured_cols,
        ):
            st.stop()
        partial_path = cache_path.replace(".pkl", "_partial.pkl")
        # Apply filters before processing so only relevant rows are translated
        df_to_process = df
        if selected_segments:
            df_to_process = df_to_process[df_to_process[location_col].isin(selected_segments)]
        for col, vals in addl_filters.items():
            df_to_process = df_to_process[df_to_process[col].isin(vals)]

        if analysis_mode != "Structured Data Only":
            with st.spinner("Processing free-text responses..."):
                processed_subset = process_free_text(
                    df_to_process,
                    free_text_cols,
                    cache_path,
                    batch_size=batch_size,
                )

            # Ensure new columns exist in the main DataFrame so update() doesn't
            # drop the AI-generated results such as Categories or reasoning.
            for col in [
                "Concatenated",
                "Translated",
                "Language",
                "Categories",
                "CategoryReasoning",
                "OriginalCategories",
                "OriginalCategoryReasoning",
                "ModelTokens",
                "FinishReason",
            ]:
                if col not in df.columns:
                    df[col] = "" if col != "ModelTokens" else 0

            # Merge processed rows back into the main dataframe so previously
            # unprocessed English comments are retained for later analysis.
            for col in processed_subset.columns:
                df.loc[processed_subset.index, col] = processed_subset[col]

            st.success("Processing complete")

            if "ModelTokens" in df.columns:
                st.session_state["total_tokens"] = int(df["ModelTokens"].sum())

            if show_comments:
                processed_subset = review_translations(processed_subset, user_id_col)
            # Persist edits back into the full DataFrame and cache
            for col in processed_subset.columns:
                df.loc[processed_subset.index, col] = processed_subset[col]
            st.session_state["processed_df"] = df
            df.to_pickle(cache_path)
            if os.path.exists(partial_path):
                os.remove(partial_path)
        else:
            st.session_state["processed_df"] = df
            if "ModelTokens" in df.columns:
                st.session_state["total_tokens"] = int(df["ModelTokens"].sum())

    processed_df = st.session_state.get("processed_df")
    if processed_df is not None:
        nps_col = next(
            (
                c
                for c in structured_cols
                if "nps" in c.lower() or "how likely are you to recommend" in str(c).lower()
            ),
            None,
        )

        analysis_df = processed_df
        if selected_segments:
            analysis_df = analysis_df[analysis_df[location_col].isin(selected_segments)]
        for col, vals in addl_filters.items():
            analysis_df = analysis_df[analysis_df[col].isin(vals)]

        display_summary(analysis_df, nps_col)

        st.subheader("Download Enriched Dataset")
        download_link(
            processed_df,
            "enriched_results.csv",
            "Download Enriched Dataset CSV",
            help="Save the processed data including translations and categories.",
        )
        export_full_excel(
            processed_df,
            "enriched_results.xlsx",
            "Download Enriched Dataset Excel",
            help="Save the processed data as an Excel workbook.",
        )

        if analysis_mode != "Free Text Only":
            st.subheader("Structured Data Analysis")
            zip_entries: list[tuple[str, bytes]] = []
            pdf_pivots: dict[str, pd.DataFrame] = {}

            # Identify multi-select question groups by numeric prefix
            pattern = re.compile(r"^(\d+)[\.:]")
            groups: dict[str, list[str]] = {}
            for col in structured_cols:
                m = pattern.match(str(col))
                if m and m.group(1) != "5":
                    groups.setdefault(m.group(1), []).append(col)

            # Filter groups to only binary multi-select columns
            groups = {
                k: v
                for k, v in groups.items()
                if len(v) > 1
                and all(
                    len(
                        {
                            str(x).strip().lower()
                            for x in analysis_df[c].dropna().unique()
                            if str(x).strip() != ""
                        }
                    )
                    <= 2
                    for c in v
                )
            }

            processed: set[str] = set()

            ease_cols = [c for c in structured_cols if re.match(r"^(6|14)[\.:]", str(c))]
            if ease_cols:
                pivot = rating_pivot(analysis_df, ease_cols, order=RATING_ORDER_EASE)
                question_text = "Of the following areas, please rate how easy they were to use"
                st.write(f"### {question_text}")
                st.dataframe(pivot)
                chart_buf = stacked_bar_chart(pivot, "Ease of Use by Area", order=RATING_ORDER_EASE)
                c1, c2 = st.columns(2)
                with c1:
                    download_link(
                        pivot,
                        "pivot_ease.csv",
                        "Download Ease Question CSV",
                        help="Download the pivot table as a CSV file.",
                    )
                with c2:
                    export_excel(
                        pivot,
                        "pivot_ease.xlsx",
                        "Download Ease Question Excel",
                        help="Download the pivot table as an Excel file.",
                    )
                csv_bytes = pivot.to_csv(index=False).encode("utf-8")
                safe = safe_name("question_ease")
                zip_entries.append((f"{safe}/table.csv", csv_bytes))
                zip_entries.append((f"{safe}/chart.png", chart_buf.getvalue()))
                pdf_pivots[question_text] = pivot
                processed.update(ease_cols)

            content_rating_cols = [c for c in structured_cols if str(c).startswith("8.")]
            if content_rating_cols:
                pivot = rating_pivot(analysis_df, content_rating_cols, order=CONTENT_RATING_ORDER)
                question_text = "Please rate the following about our content"
                st.write(f"### {question_text}")
                st.dataframe(pivot)
                chart_buf = stacked_bar_chart(pivot, "Content Ratings", order=CONTENT_RATING_ORDER)
                c1, c2 = st.columns(2)
                with c1:
                    download_link(
                        pivot,
                        "pivot_q8.csv",
                        "Download Question 8 CSV",
                        help="Download the pivot table as a CSV file.",
                    )
                with c2:
                    export_excel(
                        pivot,
                        "pivot_q8.xlsx",
                        "Download Question 8 Excel",
                        help="Download the pivot table as an Excel file.",
                    )
                csv_bytes = pivot.to_csv(index=False).encode("utf-8")
                safe = safe_name("question_8")
                zip_entries.append((f"{safe}/table.csv", csv_bytes))
                zip_entries.append((f"{safe}/chart.png", chart_buf.getvalue()))
                pdf_pivots[question_text] = pivot
                processed.update(content_rating_cols)

            satisfaction_cols = [
                c
                for c in structured_cols
                if re.match(r"^12(\.\d+)?[\.:]", str(c))
            ]
            if satisfaction_cols:
                pivot = combined_rating_pivot(
                    analysis_df, satisfaction_cols, order=SATISFACTION_ORDER
                )
                question_text = "How satisfied were you with the materials you created?"
                st.write(f"### {question_text}")
                st.dataframe(pivot)
                chart_buf = bar_chart(
                    pivot, "Satisfaction with Created Materials", order=SATISFACTION_ORDER
                )
                c1, c2 = st.columns(2)
                with c1:
                    download_link(
                        pivot,
                        "pivot_q12.csv",
                        "Download Question 12 CSV",
                        help="Download the pivot table as a CSV file.",
                    )
                with c2:
                    export_excel(
                        pivot,
                        "pivot_q12.xlsx",
                        "Download Question 12 Excel",
                        help="Download the pivot table as an Excel file.",
                    )
                csv_bytes = pivot.to_csv(index=False).encode("utf-8")
                safe = safe_name("question_12")
                zip_entries.append((f"{safe}/table.csv", csv_bytes))
                zip_entries.append((f"{safe}/chart.png", chart_buf.getvalue()))
                pdf_pivots[question_text] = pivot
                processed.update(satisfaction_cols)

            importance_cols = [c for c in structured_cols if re.match(r"^21[\.:]", str(c))]
            if importance_cols:
                vals = (
                    analysis_df[importance_cols]
                    .stack()
                    .dropna()
                    .astype(str)
                    .str.strip()
                    .str.replace(r"\.0$", "", regex=True)
                )
                if not vals.empty and vals.str.fullmatch(r"[1-5]").all():
                    imp_order = IMPORTANCE_ORDER_NUMERIC
                else:
                    imp_order = IMPORTANCE_ORDER
                pivot = rating_pivot(analysis_df, importance_cols, order=imp_order)
                question_text = "Tell us how important the following are to you"
                st.write(f"### {question_text}")
                st.dataframe(pivot)
                chart_buf = stacked_bar_chart(pivot, "Importance Ratings", order=imp_order)
                c1, c2 = st.columns(2)
                with c1:
                    download_link(
                        pivot,
                        "pivot_q21.csv",
                        "Download Question 21 CSV",
                        help="Download the pivot table as a CSV file.",
                    )
                with c2:
                    export_excel(
                        pivot,
                        "pivot_q21.xlsx",
                        "Download Question 21 Excel",
                        help="Download the pivot table as an Excel file.",
                    )
                csv_bytes = pivot.to_csv(index=False).encode("utf-8")
                safe = safe_name("question_21")
                zip_entries.append((f"{safe}/table.csv", csv_bytes))
                zip_entries.append((f"{safe}/chart.png", chart_buf.getvalue()))
                pdf_pivots[question_text] = pivot

            for prefix, cols in groups.items():
                question = MULTISELECT_QUESTION_TEXTS.get(prefix, f"Question {prefix}")
                pivot = multiselect_pivot(analysis_df, cols)
                st.write(f"### {question}")
                st.dataframe(pivot)
                chart_buf = bar_chart(pivot, f"{question} Responses")
                c1, c2 = st.columns(2)
                with c1:
                    download_link(
                        pivot,
                        f"pivot_q{prefix}.csv",
                        f"Download Question {prefix} CSV",
                        help="Download the pivot table as a CSV file.",
                    )
                with c2:
                    export_excel(
                        pivot,
                        f"pivot_q{prefix}.xlsx",
                        f"Download Question {prefix} Excel",
                        help="Download the pivot table as an Excel file.",
                    )
                csv_bytes = pivot.to_csv(index=False).encode("utf-8")
                safe = safe_name(f"question_{prefix}")
                zip_entries.append((f"{safe}/table.csv", csv_bytes))
                zip_entries.append((f"{safe}/chart.png", chart_buf.getvalue()))
                pdf_pivots[question] = pivot
                processed.update(cols)

            for col in structured_cols:
                if col in processed:
                    continue
                drop_na = "what percentage of your membership" in str(col).lower()
                pivot = generate_pivot(analysis_df, col, exclude_na=drop_na)
                question_text = str(col)
                st.write(f"### {question_text}")
                st.dataframe(pivot)
                chart_buf = bar_chart(pivot, f"{col} Responses")
                c1, c2 = st.columns(2)
                with c1:
                    download_link(
                        pivot,
                        f"pivot_{col}.csv",
                        f"Download {col} CSV",
                        help="Download the pivot table as a CSV file.",
                    )
                with c2:
                    export_excel(
                        pivot,
                        f"pivot_{col}.xlsx",
                        f"Download {col} Excel",
                        help="Download the pivot table as an Excel file.",
                    )
                csv_bytes = pivot.to_csv(index=False).encode("utf-8")
                safe = safe_name(col)
                zip_entries.append((f"{safe}/table.csv", csv_bytes))
                zip_entries.append((f"{safe}/chart.png", chart_buf.getvalue()))
                pdf_pivots[question_text] = pivot

            if zip_entries:
                zip_buf = BytesIO()
                with zipfile.ZipFile(zip_buf, "w") as zipf:
                    for name, data in zip_entries:
                        zipf.writestr(name, data)
                zip_buf.seek(0)
                st.download_button(
                    "Download All Charts/Tables",
                    zip_buf,
                    "all_pivots.zip",
                    help="Download every pivot table CSV and chart PNG at once.",
                    key=unique_key("all_pivots_zip"),
                )
            if pdf_pivots:
                pdf_buf_all = save_pdf("NPS Survey Charts and Tables", pdf_pivots)
                st.session_state["pdf_pivots"] = pdf_pivots
                st.download_button(
                    "Download Everything PDF",
                    pdf_buf_all,
                    "all_charts_tables.pdf",
                    help="Download all analysis results as a single PDF.",
                    key=unique_key("all_pivots_pdf"),
                )
                st.download_button(
                    "Download all graphs/tables on a PDF",
                    pdf_buf_all,
                    "graphs_tables.pdf",
                    help="Download all graphs and pivot tables in a single PDF.",
                    key=unique_key("graphs_tables_pdf"),
                )
                pdf_buf_charts = save_pdf(
                    "NPS Survey Charts",
                    pdf_pivots,
                    include_charts=True,
                    include_tables=False,
                )
                st.download_button(
                    "Download Charts PDF",
                    pdf_buf_charts,
                    "all_charts.pdf",
                    help="Download every chart with its question text.",
                    key=unique_key("all_charts_pdf"),
                )
                pdf_tables = save_pdf(
                    "NPS Survey Tables", pdf_pivots, include_charts=False
                )
                st.download_button(
                    "Download Tables PDF",
                    pdf_tables,
                    "all_tables.pdf",
                    help="Download every pivot table with its question text.",
                    key=unique_key("all_tables_pdf"),
                )

        if analysis_mode != "Structured Data Only" and show_comments:
            st.subheader("Categorized Comments")
            display_cols = [
                user_id_col,
                location_col,
                "Concatenated",
                "Translated",
                "Language",
                "OriginalCategories",
                "Categories",
                "ModelTokens",
                "FinishReason",
                "Flagged",
            ]
            st.dataframe(analysis_df[display_cols])
            if st.button(
                "Spot-check 5 Random Comments", help="Preview a random sample to verify AI results."
            ):
                sample = analysis_df.sample(min(5, len(analysis_df)))
                for _, row in sample.iterrows():
                    st.write(f"**User {row[user_id_col]}** - {row['Categories']}")
                    st.write(row["Translated"])

            if st.button("Generate Report", help="Create a narrative summary of all findings."):
                segments_to_process = selected_segments if selected_segments else [None]
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, "w") as zipf:
                    for segment in segments_to_process:
                        seg_df = df if segment is None else df[df[location_col] == segment]
                        for col, vals in addl_filters.items():
                            seg_df = seg_df[seg_df[col].isin(vals)]
                        if seg_df.empty:
                            continue
                        segment_title = segment if segment is not None else "All"

                        st.markdown(f"## KPIs for {segment_title}")
                        st.metric("Rows in Segment", len(seg_df))
                        seg_tokens = None
                        seg_cost = None
                        if "ModelTokens" in seg_df.columns:
                            seg_tokens = int(seg_df["ModelTokens"].sum())
                            st.metric("Total Model Tokens", seg_tokens)
                            if TOKEN_COST_PER_1K:
                                seg_cost = seg_tokens / 1000 * TOKEN_COST_PER_1K
                                st.metric("Estimated Cost", f"${seg_cost:.2f}")
                        nps_pivot, cat_pivot, (pos, neg), nps_score = compute_kpis(seg_df, nps_col)
                        if nps_score is not None:
                            st.metric("NPS Score", nps_score)
                        if not nps_pivot.empty:
                            st.dataframe(nps_pivot)
                            bar_chart(nps_pivot, f"{segment_title} NPS Distribution")
                        st.dataframe(cat_pivot)
                        cat_chart = cat_pivot[cat_pivot["Category"] != "Uncategorized"]
                        if not cat_chart.empty:
                            bar_chart(cat_chart, f"{segment_title} Category Frequency")
                        st.metric("Positive/Negative Ratio", f"{pos}:{neg}")

                        report_text = generate_report(
                            seg_df[
                                [user_id_col, location_col, "Translated", "Categories", "Flagged"]
                            ]
                        )
                        if not report_text:
                            continue
                        st.markdown(f"## Report for {segment_title}")
                        if seg_tokens is not None:
                            seg_cost_disp = seg_cost if seg_cost is not None else 0
                            st.markdown(
                                f"**Total tokens used:** {seg_tokens} (estimated cost ${seg_cost_disp:.2f})"
                            )
                        st.markdown(report_text)
                        pivot_dict = {
                            col: generate_pivot(
                                seg_df,
                                col,
                                exclude_na="what percentage of your membership" in str(col).lower(),
                            )
                            for col in structured_cols
                        }
                        pivot_dict["Category Frequency"] = cat_pivot
                        if not nps_pivot.empty:
                            pivot_dict["NPS Distribution"] = nps_pivot
                        token_prefix = ""
                        if seg_tokens is not None:
                            seg_cost_calc = seg_cost if seg_cost is not None else 0
                            token_prefix = (
                                f"Total tokens used: {seg_tokens} (estimated cost ${seg_cost_calc:.2f})\n\n"
                            )
                        docx_file = save_docx(token_prefix + report_text, pivot_dict)
                        pdf_file = save_pdf(token_prefix + report_text, pivot_dict)
                        pdf_charts = save_pdf(
                            token_prefix + report_text, pivot_dict, include_charts=True, include_tables=False
                        )
                        pdf_tables = save_pdf(
                            token_prefix + report_text, pivot_dict, include_charts=False
                        )
                        if len(selected_segments) <= 1:
                            st.download_button(
                                "Download DOCX",
                                docx_file,
                                f"{segment_title}_report.docx",
                                help="Save the report as a Word document.",
                                key=unique_key(f"{segment_title}_docx"),
                            )
                            st.download_button(
                                "Download Everything PDF",
                                pdf_file,
                                f"{segment_title}_report.pdf",
                                help="Download the full report with tables and charts as a PDF.",
                                key=unique_key(f"{segment_title}_pdf"),
                            )
                            st.download_button(
                                "Download Charts PDF",
                                pdf_charts,
                                f"{segment_title}_charts.pdf",
                                help="PDF containing each chart with question text only.",
                                key=unique_key(f"{segment_title}_charts_pdf"),
                            )
                            st.download_button(
                                "Download Tables PDF",
                                pdf_tables,
                                f"{segment_title}_tables.pdf",
                                help="PDF containing each pivot table with question text only.",
                                key=unique_key(f"{segment_title}_tables_pdf"),
                            )
                        else:
                            zipf.writestr(f"{segment_title}_report.docx", docx_file.getvalue())
                            zipf.writestr(f"{segment_title}_report.pdf", pdf_file.getvalue())
                            zipf.writestr(f"{segment_title}_charts.pdf", pdf_charts.getvalue())
                            zipf.writestr(
                                f"{segment_title}_tables.pdf", pdf_tables.getvalue()
                            )
                if len(selected_segments) > 1:
                    zip_buffer.seek(0)
                    st.download_button(
                        "Download Reports ZIP",
                        zip_buffer,
                        "reports.zip",
                        help="Download all segment reports as a ZIP file.",
                        key=unique_key("reports_zip"),
                    )
else:
    st.info("Upload a survey file to begin.")
